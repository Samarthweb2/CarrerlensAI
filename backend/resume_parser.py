import os
import re
from typing import Optional, List, Dict, Any
# pyrefly: ignore [missing-import]
import fitz  # PyMuPDF
# pyrefly: ignore [missing-import]
import docx

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extracts plain text from a PDF file using PyMuPDF (fitz).
    """
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        doc.close()
    except Exception as e:
        print(f"Error extracting text from PDF {pdf_path}: {e}")
    return text.strip()

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extracts plain text from a DOCX file using python-docx.
    """
    text_list = []
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text_list.append(para.text.strip())
        # Also extract table cells if any exist
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_list.append(" | ".join(row_text))
    except Exception as e:
        print(f"Error extracting text from DOCX {docx_path}: {e}")
    return "\n".join(text_list).strip()

def extract_text(file_path: str) -> str:
    """
    Extracts plain text from either a PDF or DOCX file.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    return ""

def extract_email(text: str) -> Optional[str]:
    """
    Extracts email address from text using regex.
    """
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    return email_match.group(0) if email_match else None

def extract_phone(text: str) -> Optional[str]:
    """
    Extracts phone number from text using regex.
    """
    phone_match = re.search(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    return phone_match.group(0) if phone_match else None

def extract_links(text: str) -> Dict[str, Optional[str]]:
    """
    Extracts GitHub and LinkedIn handles from text.
    """
    github_match = re.search(r'(github\.com/[\w\.-]+)', text, re.IGNORECASE)
    linkedin_match = re.search(r'(linkedin\.com/in/[\w\.-]+)', text, re.IGNORECASE)
    
    return {
        "github": github_match.group(0) if github_match else None,
        "linkedin": linkedin_match.group(0) if linkedin_match else None
    }

def extract_skills_keywords(text: str) -> List[str]:
    """
    Matches text against a comprehensive list of tech, data science, GenAI, and MLOps skills.
    """
    title_case_mapping = {
        "python": "Python", "sql": "SQL", "power bi": "Power BI", "react": "React",
        "fastapi": "FastAPI", "machine learning": "Machine Learning", "pandas": "Pandas",
        "numpy": "NumPy", "docker": "Docker", "aws": "AWS", "kubernetes": "Kubernetes",
        "ci/cd": "CI/CD", "javascript": "JavaScript", "html": "HTML", "css": "CSS",
        "git": "Git", "java": "Java", "c++": "C++", "tableau": "Tableau", "excel": "Excel",
        "spark": "Spark", "pyspark": "PySpark", "scikit-learn": "Scikit-Learn", "scikit learn": "Scikit-Learn",
        "hadoop": "Hadoop", "nlp": "NLP", "deep learning": "Deep Learning", "pytorch": "PyTorch",
        "tensorflow": "TensorFlow", "agile": "Agile", "scrum": "Scrum",
        "rag": "RAG", "agentic ai": "Agentic AI", "llm": "LLMs", "llms": "LLMs",
        "multi-agent": "Multi-Agent Systems", "multi-agent systems": "Multi-Agent Systems",
        "langchain": "LangChain", "langgraph": "LangGraph", "llamaindex": "LlamaIndex",
        "ragas": "RAGAS", "langsmith": "LangSmith", "bigquery": "BigQuery", "airflow": "Airflow",
        "flask": "Flask", "rest apis": "REST APIs", "rest api": "REST API", "sqlite": "SQLite",
        "data science": "Data Science", "data scientist": "Data Science", "data analysis": "Data Analysis",
        "statistical modeling": "Statistical Modeling", "statistics": "Statistics",
        "time series": "Time Series", "marketing mix modeling": "Marketing Mix Modeling",
        "recommender systems": "Recommender Systems", "multimodal ai": "Multimodal AI",
        "bayesian": "Bayesian Modeling", "azure ai studio": "Azure AI Studio",
        "dbt": "dbt", "snowflake": "Snowflake", "redis": "Redis", "postgresql": "PostgreSQL",
        "temporal": "Temporal", "transformers": "Transformers"
    }
    
    extracted = []
    text_lower = text.lower()
    for skill_key, display_name in title_case_mapping.items():
        pattern = r'(?:\b|(?<=\W))' + re.escape(skill_key) + r'(?:\b|(?=\W))'
        if re.search(pattern, text_lower):
            extracted.append(display_name)
            
    # Deduplicate while preserving cased names
    return list(dict.fromkeys(extracted))

def extract_section_content(text: str, section_name: str) -> List[str]:
    """
    Heuristically extracts lines located under specific section headers.
    """
    lines = text.split('\n')
    section_headers = {
        "summary": ["summary", "professional summary", "executive summary", "profile", "about me", "career summary", "objective"],
        "education": ["education", "academic background", "studies", "qualification", "qualifications", "academic credentials"],
        "experience": ["experience", "employment history", "work experience", "professional experience", "internships", "employment"],
        "projects": ["projects", "personal projects", "academic projects", "key projects", "development projects"],
        "skills": ["skills", "technical skills", "key skills", "technologies", "expertise", "competencies"],
        "certifications": ["certifications", "licenses", "courses", "credentials"]
    }
    
    targets = section_headers.get(section_name.lower(), [])
    if not targets:
        return []
        
    start_idx = -1
    for i, line in enumerate(lines):
        clean_line = line.strip().lower().rstrip(':')
        if clean_line in targets or any(clean_line == t for t in targets):
            start_idx = i
            break
            
    if start_idx == -1:
        # Fuzzy match
        for i, line in enumerate(lines):
            clean_line = line.strip().lower().rstrip(':')
            if any(t in clean_line for t in targets) and len(clean_line.split()) <= 3:
                start_idx = i
                break
                
    if start_idx == -1:
        return []
        
    content_lines = []
    all_headers = []
    for h_list in section_headers.values():
        all_headers.extend(h_list)
        
    for line in lines[start_idx + 1:]:
        clean_line = line.strip().lower().rstrip(':')
        if any(clean_line == h for h in all_headers) or (any(h in clean_line for h in all_headers) and len(clean_line.split()) <= 3):
            break
        if line.strip():
            content_lines.append(line.strip())
            
    return content_lines

def extract_name(text: str) -> Optional[str]:
    """
    Heuristically extracts candidate name from the top of the resume text.
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]:
        if "@" in line or any(p in line.lower() for p in ["+", "phone", "email", "github", "linkedin", "resume", "cv"]):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(part[0].isupper() for part in words if part.isalpha()):
            return line
    return None

def parse_resume_to_json(file_path: str, filename: str) -> Dict[str, Any]:
    """
    Main parser coordinator. Extracts text and segments it into structured JSON fields.
    """
    raw_text = extract_text(file_path)
    
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    links = extract_links(raw_text)
    
    summary_lines = extract_section_content(raw_text, "summary")
    summary = " ".join(summary_lines) if summary_lines else ""
    if not summary:
        for line in [l.strip() for l in raw_text.split('\n') if l.strip()][:15]:
            if len(line) > 50 and not any(k in line.lower() for k in ['email', '@', 'phone', 'github', 'linkedin', 'http']):
                summary = line
                break
    
    education = extract_section_content(raw_text, "education")
    experience = extract_section_content(raw_text, "experience")
    projects = extract_section_content(raw_text, "projects")
    skills = extract_skills_keywords(raw_text)
    certifications = extract_section_content(raw_text, "certifications")
    
    extracted_name = extract_name(raw_text)
    
    return {
        "fileId": str(os.path.basename(file_path).split('_')[0]),
        "fileName": filename,
        "name": extracted_name or "Candidate Name",
        "email": email,
        "phone": phone,
        "text": raw_text,
        "summary": summary,
        "education": education,
        "experience": experience,
        "projects": projects,
        "skills": skills,
        "certifications": certifications,
        "links": links
    }
